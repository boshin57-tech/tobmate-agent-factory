from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


SOURCE_DIR = Path("workspace/source")
INDEX_DIR = Path("workspace/index")
CHUNK_SIZE = 1800
OVERLAP = 250


def normalize(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text:
            parts.append(text)

    for table_no, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))

        if rows:
            parts.append(f"[TABLE {table_no}]\n" + "\n".join(rows))

    return normalize("\n\n".join(parts))


def split_text(text: str):
    start = 0
    number = 1

    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))

        if end < len(text):
            breaks = [
                text.rfind("\n\n", start, end),
                text.rfind("\n", start, end),
                text.rfind(". ", start, end),
            ]
            best = max(breaks)

            if best > start + CHUNK_SIZE // 2:
                end = best + 1

        chunk = normalize(text[start:end])

        if chunk:
            yield number, chunk
            number += 1

        if end >= len(text):
            break

        start = max(end - OVERLAP, start + 1)


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()

    with path.open("rb") as f:
        while True:
            block = f.read(1024 * 1024)
            if not block:
                break
            digest.update(block)

    return digest.hexdigest()


def main() -> None:
    if not SOURCE_DIR.exists():
        raise SystemExit(f"Source folder missing: {SOURCE_DIR}")

    INDEX_DIR.mkdir(parents=True, exist_ok=True)

    files = sorted(SOURCE_DIR.rglob("*.docx"))
    chunk_file = INDEX_DIR / "source_chunks.jsonl"
    manifest_file = INDEX_DIR / "source_manifest.json"
    error_file = INDEX_DIR / "source_errors.json"

    manifest_files = []
    errors = []
    total_chunks = 0
    total_characters = 0

    with chunk_file.open("w", encoding="utf-8") as out:
        for doc_no, path in enumerate(files, start=1):
            relative = path.relative_to(SOURCE_DIR)

            try:
                text = extract_docx(path)
                chunks = list(split_text(text))

                document_id = hashlib.sha1(
                    str(relative).encode("utf-8")
                ).hexdigest()[:16]

                for chunk_no, chunk in chunks:
                    record = {
                        "chunk_id": f"{document_id}-{chunk_no:04d}",
                        "document_id": document_id,
                        "document_number": doc_no,
                        "chunk_number": chunk_no,
                        "source_path": str(relative),
                        "file_name": path.name,
                        "category": relative.parts[0] if relative.parts else "source",
                        "text": chunk,
                        "character_count": len(chunk),
                    }

                    out.write(json.dumps(record, ensure_ascii=False) + "\n")

                total_chunks += len(chunks)
                total_characters += len(text)

                manifest_files.append(
                    {
                        "document_id": document_id,
                        "source_path": str(relative),
                        "file_name": path.name,
                        "size_bytes": path.stat().st_size,
                        "sha256": file_hash(path),
                        "character_count": len(text),
                        "chunk_count": len(chunks),
                    }
                )

                print(
                    f"[{doc_no:03d}/{len(files):03d}] "
                    f"{relative} -> {len(chunks)} chunks"
                )

            except Exception as exc:
                errors.append(
                    {
                        "source_path": str(relative),
                        "error": f"{type(exc).__name__}: {exc}",
                    }
                )
                print(f"[ERROR] {relative}: {exc}")

    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "document_count": len(files),
        "indexed_document_count": len(manifest_files),
        "error_count": len(errors),
        "total_chunks": total_chunks,
        "total_characters": total_characters,
        "files": manifest_files,
    }

    manifest_file.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    error_file.write_text(
        json.dumps(errors, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print()
    print("=== SOURCE INDEX COMPLETE ===")
    print("Documents discovered :", len(files))
    print("Documents indexed    :", len(manifest_files))
    print("Chunks generated     :", total_chunks)
    print("Errors               :", len(errors))
    print("Index file           :", chunk_file)


if __name__ == "__main__":
    main()
